from docx import Document,shared
from PIL import Image
from gen_docx import gen_docx
from pillow import change_dpi

def map_colors_to_document(docx_path, image_path):
    # 打开docx文档
    doc = Document(docx_path)

    # 打开图片并获取像素数据
    image = Image.open(image_path)
    pixels = list(image.getdata())

    run_index = 0
    new_doc = Document()
    # 获取文档的第一节（默认有一个节）
    section = new_doc.sections[0]

    # 设置页面边距（上下左右均为2cm）
    section.left_margin = shared.Cm(2)
    section.right_margin = shared.Cm(2)
    section.top_margin = shared.Cm(2)
    section.bottom_margin = shared.Cm(2)
    # 遍历文档中的段落和字体
    for paragraph in doc.paragraphs:
        new_paragraph = new_doc.add_paragraph()
        for run in paragraph.runs:
            for char in run.text:
                if run_index < len(pixels):
                    new_run = new_paragraph.add_run(char)
                    new_run.font.color.rgb = shared.RGBColor(*pixels[run_index])
                    run_index += 1
                else:
                    break
        # 设置字体大小为5.5
        new_paragraph.style.font.size = shared.Pt(5.5)

        # 设置行间距为5.5磅
        new_paragraph.paragraph_format.line_spacing = shared.Pt(5.5)

    # 保存修改后的文档
    modified_docx_path = "modified_" + docx_path
    new_doc.save(modified_docx_path)
    print(f"文档已保存为: {modified_docx_path}")

if __name__ == "__main__":
    docx_path = gen_docx()
    image_path = change_dpi()
    map_colors_to_document(docx_path, image_path)
