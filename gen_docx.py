from docx import Document, shared

def gen_docx():
    new_docx_name = "ikun.docx"
    # 创建一个新的 docx 文档
    document = Document()
    # 获取文档的第一节（默认有一个节）
    section = document.sections[0]

    # 设置页面边距（上下左右均为2cm）
    section.left_margin = shared.Cm(2)
    section.right_margin = shared.Cm(2)
    section.top_margin = shared.Cm(2)
    section.bottom_margin = shared.Cm(2)

    # 添加一个段落，并设置其对齐方式为居中
    paragraph = document.add_paragraph()
    paragraph.alignment = 1

    # 重复添加文本“蔡徐坤”直至达到目标长度
    text = ""
    while len(text) < 90 * 120:
        text += "蔡徐坤"

    # 将文本添加到段落中
    paragraph.add_run(text)
    # 设置字体大小为5.5
    paragraph.style.font.size = shared.Pt(5.5)

    # 设置行间距为5.5磅
    paragraph.paragraph_format.line_spacing = shared.Pt(5.5)

    # 保存文档
    document.save(new_docx_name)
    return new_docx_name